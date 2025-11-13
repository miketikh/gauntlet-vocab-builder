"""
Text Extraction Service
Handles extraction of text from various file formats (PDF, DOCX, TXT)
"""
from io import BytesIO
from typing import Optional
import logging

# PDF extraction
try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

# DOCX extraction
try:
    from docx import Document
except ImportError:
    Document = None

# File type detection
try:
    import magic
except ImportError:
    magic = None


logger = logging.getLogger(__name__)


class TextExtractionError(Exception):
    """Raised when text extraction fails"""
    pass


class UnsupportedFileTypeError(TextExtractionError):
    """Raised when file type is not supported"""
    pass


def extract_from_pdf(content: bytes) -> str:
    """
    Extract text from PDF file content

    Args:
        content: PDF file bytes

    Returns:
        Extracted text as string

    Raises:
        TextExtractionError: If PDF extraction fails
    """
    if PdfReader is None:
        raise TextExtractionError("pypdf library not installed. Run: pip install pypdf")

    try:
        pdf_file = BytesIO(content)
        reader = PdfReader(pdf_file)

        text_parts = []
        for page in reader.pages:
            try:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            except Exception as e:
                logger.warning(f"Failed to extract text from PDF page: {e}")
                continue

        if not text_parts:
            raise TextExtractionError("No text could be extracted from PDF")

        return "\n\n".join(text_parts)

    except Exception as e:
        raise TextExtractionError(f"Failed to extract text from PDF: {str(e)}")


def extract_from_docx(content: bytes) -> str:
    """
    Extract text from DOCX file content

    Args:
        content: DOCX file bytes

    Returns:
        Extracted text as string

    Raises:
        TextExtractionError: If DOCX extraction fails
    """
    if Document is None:
        raise TextExtractionError("python-docx library not installed. Run: pip install python-docx")

    try:
        docx_file = BytesIO(content)
        doc = Document(docx_file)

        # Extract paragraphs
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

        # Extract text from tables
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text.strip()]
                if row_text:
                    table_text.append(" ".join(row_text))

        all_text = paragraphs + table_text

        if not all_text:
            raise TextExtractionError("No text could be extracted from DOCX")

        return "\n\n".join(all_text)

    except Exception as e:
        raise TextExtractionError(f"Failed to extract text from DOCX: {str(e)}")


def extract_from_txt(content: bytes) -> str:
    """
    Extract text from plain text file content
    Handles various text encodings gracefully

    Args:
        content: Text file bytes

    Returns:
        Decoded text as string

    Raises:
        TextExtractionError: If text decoding fails
    """
    # Try common encodings in order
    encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']

    for encoding in encodings:
        try:
            text = content.decode(encoding)
            # Check if text is not empty after stripping
            if text.strip():
                return text
        except (UnicodeDecodeError, AttributeError):
            continue

    # If all encodings fail, try with error handling
    try:
        text = content.decode('utf-8', errors='replace')
        if text.strip():
            logger.warning("Text decoded with replacement characters for invalid bytes")
            return text
    except Exception as e:
        raise TextExtractionError(f"Failed to decode text file: {str(e)}")

    raise TextExtractionError("Could not decode text file with any supported encoding")


def detect_file_type(content: bytes, filename: Optional[str] = None) -> str:
    """
    Detect file type from content or filename

    Args:
        content: File bytes
        filename: Optional filename with extension

    Returns:
        File type: 'pdf', 'docx', 'txt', or 'unknown'
    """
    # First try magic library if available
    if magic is not None:
        try:
            mime_type = magic.from_buffer(content, mime=True)
            if mime_type == 'application/pdf':
                return 'pdf'
            elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                               'application/msword']:
                return 'docx'
            elif mime_type.startswith('text/'):
                return 'txt'
        except Exception as e:
            logger.warning(f"Magic library failed to detect file type: {e}")

    # Fallback to filename extension
    if filename:
        filename_lower = filename.lower()
        if filename_lower.endswith('.pdf'):
            return 'pdf'
        elif filename_lower.endswith('.docx'):
            return 'docx'
        elif filename_lower.endswith('.txt'):
            return 'txt'

    # Fallback to content inspection
    # Check for PDF signature
    if content.startswith(b'%PDF'):
        return 'pdf'

    # Check for DOCX (ZIP) signature
    if content.startswith(b'PK\x03\x04') and b'word/' in content[:1000]:
        return 'docx'

    # Default to text
    return 'txt'


def extract_text_from_file(
    file_content: bytes,
    file_type: Optional[str] = None,
    filename: Optional[str] = None
) -> str:
    """
    Main text extraction function - extracts text from various file formats

    Args:
        file_content: File bytes to extract text from
        file_type: Optional file type hint ('pdf', 'docx', 'txt')
        filename: Optional filename for type detection

    Returns:
        Extracted and cleaned text

    Raises:
        TextExtractionError: If extraction fails
        UnsupportedFileTypeError: If file type is not supported
    """
    # Validate input
    if not file_content:
        raise TextExtractionError("File content is empty")

    # Detect file type if not provided
    if not file_type:
        file_type = detect_file_type(file_content, filename)
        logger.info(f"Detected file type: {file_type}")

    # Extract text based on file type
    file_type = file_type.lower()

    if file_type == 'pdf':
        text = extract_from_pdf(file_content)
    elif file_type == 'docx':
        text = extract_from_docx(file_content)
    elif file_type == 'txt':
        text = extract_from_txt(file_content)
    else:
        raise UnsupportedFileTypeError(f"Unsupported file type: {file_type}")

    # Clean and normalize text
    text = clean_text(text)

    if not text.strip():
        raise TextExtractionError("No text content found in file")

    return text


def clean_text(text: str) -> str:
    """
    Clean and normalize extracted text

    Args:
        text: Raw extracted text

    Returns:
        Cleaned text
    """
    # Remove excessive whitespace
    lines = [line.strip() for line in text.split('\n')]

    # Remove empty lines
    lines = [line for line in lines if line]

    # Join with single newline
    text = '\n'.join(lines)

    # Replace multiple spaces with single space
    import re
    text = re.sub(r' +', ' ', text)

    return text
